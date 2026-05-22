"""
Main FFP-report DOCX writer.

Produces a Fitness-For-Purpose report mirroring the structure of the
Kandla-Samakhiali / HMEL FFP reports Athena PowerTech delivers:

  Cover page
  Executive Summary
  Abbreviations
  Table of Contents (auto-built by Word from heading styles)
  1. Introduction
  2. ILI Results
  3. Fitness-For-Purpose Analysis
  4. ILI Reports Comparison / CGR / Repair Prediction
  Disclaimer
  Annexure A — Guidelines and Formulas

The narrative for each section is in `/templates/sections/*.txt` with
`{{PLACEHOLDER}}` tokens — Athena staff edit those files directly without
touching code. Auto-populated tables and matplotlib charts (depth /
length / orientation / repair-timeline / ERF acceptance) are inserted
programmatically.

Performance budget: <5 s for a Kandla-scale (~333 features) project.
matplotlib is configured for the non-interactive Agg backend so charts
render without a display.

All inputs come from the validated pipeline outputs:
  - `Project` (config + pipeline + run metadata)
  - `MatchResult` (joint + feature matches)
  - `list[CGRResult]` (per-feature growth rates)
  - `list[FFPResult]` (per-feature FFP outcomes, one per controlling method)
  - `list[RepairPrediction]` (per-feature repair year)
  - `FlagReport` (QA findings)

The writer never reads from disk during the run beyond loading the
templates — every data point comes from these in-memory objects.
"""
from __future__ import annotations

import math
import re
import tempfile
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

# Non-interactive matplotlib backend — must be set before pyplot import.
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from src.core.repair_predictor import TRIGGER_DEPTH_80, TRIGGER_ERF_1, TRIGGER_NONE
from src.models import FFPMethod, Surface


# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------

_DEFAULT_TEMPLATES_DIR = Path(__file__).resolve().parents[2] / "templates" / "sections"
_TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x79)
_HEADING_COLOR = RGBColor(0x2E, 0x74, 0xB5)
_HEADER_FILL_HEX = "BDD7EE"


# ---------------------------------------------------------------------------
# MainReportWriter
# ---------------------------------------------------------------------------

class MainReportWriter:
    """Assemble a full FFP DOCX report from the validated pipeline outputs."""

    def __init__(self, template_dir: Path | None = None):
        self.template_dir = Path(template_dir) if template_dir else _DEFAULT_TEMPLATES_DIR

    # ------------------------------------------------------------------

    def write(
        self,
        *,
        project,
        match_result: Any = None,
        joint_alignment: Any = None,
        cgr_results: Iterable = (),
        ffp_results: Iterable = (),
        repair_predictions: Iterable = (),
        flag_report: Any = None,
        output_path: str | Path,
    ) -> Path:
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        cgr_list = list(cgr_results)
        ffp_list = list(ffp_results)
        pred_list = list(repair_predictions)

        # ----- placeholder map (one source of truth for substitutions)
        placeholders = _build_placeholders(
            project=project,
            match_result=match_result,
            joint_alignment=joint_alignment,
            cgr_results=cgr_list,
            ffp_results=ffp_list,
            repair_predictions=pred_list,
            flag_report=flag_report,
        )

        # ----- build the document
        doc = Document()
        _configure_default_styles(doc)

        with tempfile.TemporaryDirectory(prefix="ffp_charts_") as chart_dir:
            chart_dir_path = Path(chart_dir)

            self._add_cover_page(doc, project, placeholders)
            self._add_executive_summary(doc, project, placeholders,
                                        ffp_list, pred_list, flag_report)
            self._add_abbreviations(doc, placeholders)
            self._add_table_of_contents(doc)

            self._add_introduction(doc, project, placeholders)
            self._add_ili_results(doc, project, placeholders,
                                  cgr_list, chart_dir_path)
            self._add_ffp_analysis(doc, project, placeholders,
                                   ffp_list, chart_dir_path)
            self._add_cgr_repair_section(doc, project, placeholders,
                                         cgr_list, ffp_list, pred_list,
                                         chart_dir_path)

            self._add_disclaimer(doc, placeholders)
            self._add_annexure_a(doc, placeholders)

            doc.save(out_path)

        return out_path

    # ------------------------------------------------------------------
    # Section: Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc, project, ph):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Fitness-For-Purpose Report")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = _TITLE_COLOR

        doc.add_paragraph().add_run()  # spacer

        for label, val in [
            ("Pipeline", ph.get("PIPELINE_NAME", "")),
            ("Client",   ph.get("CLIENT_NAME", "")),
            ("Project",  ph.get("PROJECT_NAME", "")),
            ("Report number", ph.get("REPORT_NUMBER", "—")),
            ("Revision", ph.get("REPORT_REVISION", "00")),
            ("Date",     ph.get("REPORT_DATE", date.today().strftime("%d %b %Y"))),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{label}:  ")
            r.bold = True
            r.font.size = Pt(14)
            p.add_run(str(val)).font.size = Pt(14)

        # Revision history table — single row for now.
        doc.add_paragraph().add_run()
        rev_table = _make_table(
            doc,
            headers=["Rev", "Date", "Prepared by", "Reviewed by", "Approved by"],
            rows=[[
                ph.get("REPORT_REVISION", "00"),
                ph.get("REPORT_DATE", date.today().strftime("%d-%m-%Y")),
                ph.get("PREPARED_BY", ""),
                ph.get("REVIEWED_BY", ""),
                ph.get("APPROVED_BY", ""),
            ]],
        )
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section: Executive Summary
    # ------------------------------------------------------------------

    def _add_executive_summary(self, doc, project, ph,
                                ffp_results, predictions, flag_report):
        doc.add_heading("Executive Summary", level=1)

        n_total = int(ph.get("TOTAL_FEATURES_RUN2", 0) or 0)
        n_erf_ge_1 = int(ph.get("COUNT_ERF_GE_1", 0) or 0)
        n_d_ge_80 = int(ph.get("COUNT_DEPTH_GE_80", 0) or 0)
        n_repair = int(ph.get("COUNT_REPAIR_WITHIN_HORIZON", 0) or 0)

        for line in [
            f"Pipeline: {ph.get('PIPELINE_NAME', '')} ({ph.get('PIPELINE_LENGTH_KM', '—')} km).",
            f"Inspection comparison: {ph.get('RUN1_YEAR', '?')} vs {ph.get('RUN2_YEAR', '?')} (interval {ph.get('YEARS_BETWEEN', '?')} years).",
            f"Total metal-loss features assessed: {n_total}.",
            f"Features with ERF ≥ 1.0 today: {n_erf_ge_1}.",
            f"Features with measured depth ≥ 80% WT: {n_d_ge_80}.",
            f"Features requiring repair within the {ph.get('HORIZON_YEARS', '10')}-year horizon: {n_repair}.",
            f"Maximum ERF observed: {ph.get('MAX_ERF', '—')} (feature {ph.get('MAX_ERF_FEATURE_ID', '—')}).",
            f"Maximum depth observed: {ph.get('MAX_DEPTH_PCT', '—')} % WT (feature {ph.get('MAX_DEPTH_FEATURE_ID', '—')}).",
            f"Population P95 CGR — internal {ph.get('P95_CGR_INTERNAL', '—')} mm/yr, external {ph.get('P95_CGR_EXTERNAL', '—')} mm/yr.",
        ]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line)

        # QA summary line
        if flag_report is not None and getattr(flag_report, "summary", None):
            p = doc.add_paragraph()
            r = p.add_run("QA: ")
            r.bold = True
            p.add_run(flag_report.summary)

        # Verdict
        verdict_p = doc.add_paragraph()
        r = verdict_p.add_run("Overall verdict: ")
        r.bold = True
        if n_repair == 0 and n_erf_ge_1 == 0 and n_d_ge_80 == 0:
            verdict_p.add_run(
                f"No defects require repair within the {ph.get('HORIZON_YEARS', '10')}-year horizon. "
                "Continue operation at MAOP under the existing CP and inspection schedule."
            )
        else:
            verdict_p.add_run(
                "Defects flagged for action. See §4.3 for the response-category breakdown "
                "and §4.6 for the recommended repair list."
            )

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section: Abbreviations
    # ------------------------------------------------------------------

    def _add_abbreviations(self, doc, ph):
        doc.add_heading("Abbreviations", level=1)
        body = _load_template(self.template_dir, "abbreviations.txt", ph)
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(10)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section: Table of Contents
    # ------------------------------------------------------------------

    def _add_table_of_contents(self, doc):
        doc.add_heading("Table of Contents", level=1)
        # Insert a TOC field — Word builds the actual TOC on first open.
        p = doc.add_paragraph()
        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click and Update Field in Word to refresh."
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instr)
        run._r.append(fldChar2)
        run._r.append(placeholder)
        run._r.append(fldChar3)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section 1 — Introduction
    # ------------------------------------------------------------------

    def _add_introduction(self, doc, project, ph):
        doc.add_heading("1. Introduction", level=1)
        body = _load_template(self.template_dir, "introduction.txt", ph)
        _add_markdown_body(doc, body)

        # Pipeline details table. MAOP and design factor render as a
        # slash-separated list when there are multiple zones (e.g.
        # HMEL: "96.7 / 84.1 / 80.6 kg/cm² (by WT zone)") so the
        # single-row Pipeline-Data table doesn't silently hide
        # zone-2/zone-3 values. Full per-zone breakdown follows in
        # Table 2a below.
        pipeline = project.pipeline if project else None
        rows = [
            ["Pipeline name", ph.get("PIPELINE_NAME", "")],
            ["Client", ph.get("CLIENT_NAME", "")],
            ["Product", ph.get("PRODUCT", "")],
            ["Outside diameter", f"{ph.get('DIAMETER_MM', '—')} mm"],
            ["Length", f"{ph.get('PIPELINE_LENGTH_KM', '—')} km"],
            ["Material grade", ph.get("MATERIAL_GRADE", "")],
            ["SMYS", f"{ph.get('SMYS_MPA', '—')} MPa"],
            ["MAOP", ph.get("MAOP_DISPLAY", f"{ph.get('MAOP_KGCM2', '—')} kg/cm²")],
            ["Design factor", ph.get("DESIGN_FACTOR_DISPLAY",
                                     str(ph.get("DESIGN_FACTOR", "—")))],
            ["Service class", ph.get("SERVICE_CLASS", "")],
            ["Installation year", str(ph.get("INSTALL_YEAR", "—"))],
        ]
        _make_table(doc, headers=["Parameter", "Value"], rows=rows,
                    col_widths_in=[2.5, 4.0])

        # Table 2a — MAOP zones (full per-zone breakdown). Only emitted
        # when there are 2+ zones; for a single-zone line the headline
        # row above already carries everything the reader needs.
        # v0.3.0: heading + bound columns adapt to chainage zoning mode.
        if pipeline and pipeline.maop_zones and len(pipeline.maop_zones) >= 2:
            doc.add_paragraph()
            mode = getattr(pipeline, "maop_zoning_mode", "wt")
            if mode == "chainage":
                doc.add_heading(
                    "Table 2a — MAOP zones (by chainage)", level=3,
                )
                zone_rows = [
                    [
                        f"{z.chainage_m_min:g}" if z.chainage_m_min is not None else "—",
                        f"{z.chainage_m_max:g}" if z.chainage_m_max is not None else "—",
                        f"{z.design_factor:.2f}",
                        f"{z.maop_kgcm2:.1f}",
                    ]
                    for z in pipeline.maop_zones
                ]
                headers = ["Chainage min (m)", "Chainage max (m)",
                           "Design factor", "MAOP (kg/cm²)"]
            else:
                doc.add_heading(
                    "Table 2a — MAOP zones (by wall thickness)", level=3,
                )
                zone_rows = [
                    [
                        f"{z.wt_mm_min:g}" if z.wt_mm_min is not None else "—",
                        f"{z.wt_mm_max:g}" if z.wt_mm_max is not None else "—",
                        f"{z.design_factor:.2f}",
                        f"{z.maop_kgcm2:.1f}",
                    ]
                    for z in pipeline.maop_zones
                ]
                headers = ["WT min (mm)", "WT max (mm)",
                           "Design factor", "MAOP (kg/cm²)"]
            _make_table(
                doc, headers=headers, rows=zone_rows,
                col_widths_in=[1.5, 1.5, 1.4, 1.6],
            )
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section 2 — ILI Results
    # ------------------------------------------------------------------

    def _add_ili_results(self, doc, project, ph, cgr_list, chart_dir):
        doc.add_heading("2. ILI Results", level=1)
        body = _load_template(self.template_dir, "ili_results_discussion.txt", ph)
        _add_markdown_body(doc, body)

        # Charts: depth, length, orientation
        depth_png = chart_dir / "depth_distribution.png"
        length_png = chart_dir / "length_distribution.png"
        orient_png = chart_dir / "orientation_polar.png"

        _chart_depth_distribution(cgr_list, depth_png)
        _chart_length_distribution(cgr_list, length_png)
        _chart_orientation_polar(cgr_list, orient_png)

        for caption, png in [
            ("Figure 1. Metal-loss defect depth distribution.", depth_png),
            ("Figure 2. Metal-loss defect length distribution.", length_png),
            ("Figure 3. Defect orientation (clock position).", orient_png),
        ]:
            if png.exists():
                doc.add_picture(str(png), width=Inches(6.0))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(caption)
                r.italic = True
                r.font.size = Pt(9)

        # Shape categorisation table
        shape_rows = _shape_categorization_rows(cgr_list)
        if shape_rows:
            doc.add_heading("Shape Categorization", level=2)
            _make_table(
                doc,
                headers=["Dimension class", "Internal", "External", "Total", "%"],
                rows=shape_rows,
            )
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section 3 — FFP Analysis
    # ------------------------------------------------------------------

    def _add_ffp_analysis(self, doc, project, ph, ffp_results, chart_dir):
        doc.add_heading("3. Fitness-For-Purpose Analysis", level=1)
        body = _load_template(self.template_dir, "ffp_analysis.txt", ph)
        _add_markdown_body(doc, body)

        # ERF acceptance chart per dominant WT
        wts = sorted({round(r.wt_mm, 1) for r in ffp_results if r.wt_mm})
        for i, wt in enumerate(wts[:3]):     # show up to 3 dominant WT zones
            png = chart_dir / f"erf_chart_wt_{wt}.png"
            _chart_erf_acceptance(ffp_results, wt, png)
            if png.exists():
                doc.add_picture(str(png), width=Inches(6.0))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"Figure {4 + i}. ERF acceptance chart — WT {wt} mm.")
                r.italic = True
                r.font.size = Pt(9)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section 4 — CGR + Repair Prediction
    # ------------------------------------------------------------------

    def _add_cgr_repair_section(
        self, doc, project, ph,
        cgr_list, ffp_results, predictions, chart_dir,
    ):
        doc.add_heading(
            "4. ILI Reports Comparison, CGR Analysis, and Repair Prediction",
            level=1,
        )
        body = _load_template(self.template_dir, "cgr_methodology.txt", ph)
        _add_markdown_body(doc, body)

        # ILI comparison summary table
        doc.add_heading("Table 4 — ILI Comparison Summary", level=2)
        _make_table(
            doc,
            headers=["Quantity",
                     f"ILI {ph.get('RUN1_YEAR', 'run 1')}",
                     f"ILI {ph.get('RUN2_YEAR', 'run 2')}"],
            rows=[
                ["Total joints inspected",
                 str(ph.get("TOTAL_JOINTS_RUN1", "—")),
                 str(ph.get("TOTAL_JOINTS_RUN2", "—"))],
                ["Total metal-loss features",
                 str(ph.get("RUN1_FEATURE_COUNT", "—")),
                 str(ph.get("RUN2_FEATURE_COUNT", "—"))],
                ["Matched feature pairs",
                 str(ph.get("DEFECT_MATCH_COUNT", "—")),
                 str(ph.get("DEFECT_MATCH_COUNT", "—"))],
                ["Joint alignment rate",
                 f"{ph.get('JOINT_MATCH_RATE_PCT', '—')} %", ""],
            ],
        )

        # Upper-bound CGR table
        doc.add_heading("Table 5 — Upper-bound Corrosion Growth Rates", level=2)
        _make_table(
            doc,
            headers=["Surface", "Feature count", "P95 CGR (mm/yr)"],
            rows=[
                ["Internal", str(ph.get("COUNT_INTERNAL", "—")),
                 f"{ph.get('P95_CGR_INTERNAL', '—')}"],
                ["External", str(ph.get("COUNT_EXTERNAL", "—")),
                 f"{ph.get('P95_CGR_EXTERNAL', '—')}"],
            ],
        )

        # Top-20 by ERF
        doc.add_heading("Table 6a — Top 20 features by ERF", level=2)
        top_erf = _top_n_by_erf(ffp_results, n=20)
        if top_erf:
            _make_table(
                doc,
                headers=["#", "Feature ID", "Joint", "Chainage (m)",
                         "WT (mm)", "Depth (%WT)", "L (mm)", "Surface",
                         "Psafe (kg/cm²)", "ERF"],
                rows=top_erf,
            )

        # Top-20 by depth
        doc.add_heading("Table 6b — Top 20 features by depth", level=2)
        top_depth = _top_n_by_depth(ffp_results, n=20)
        if top_depth:
            _make_table(
                doc,
                headers=["#", "Feature ID", "Joint", "Chainage (m)",
                         "WT (mm)", "Depth (%WT)", "L (mm)", "Surface",
                         "Psafe (kg/cm²)", "ERF"],
                rows=top_depth,
            )

        # Repair timeline chart (only when something to plot)
        timeline_png = chart_dir / "repair_timeline.png"
        _chart_repair_timeline(predictions, timeline_png)
        if timeline_png.exists():
            doc.add_picture(str(timeline_png), width=Inches(6.0))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("Figure 7. Predicted-repair cumulative timeline.")
            r.italic = True
            r.font.size = Pt(9)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section: Disclaimer
    # ------------------------------------------------------------------

    def _add_disclaimer(self, doc, ph):
        doc.add_heading("Disclaimer", level=1)
        body = _load_template(self.template_dir, "disclaimer.txt", ph)
        _add_markdown_body(doc, body)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Section: Annexure A — formulas
    # ------------------------------------------------------------------

    def _add_annexure_a(self, doc, ph):
        doc.add_heading("Annexure A — Guidelines and Formulas", level=1)
        body = _load_template(self.template_dir, "annexure_a_formulas.txt", ph)
        _add_markdown_body(doc, body)


# ---------------------------------------------------------------------------
# Helpers — template loading + markdown-ish body insertion
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def _load_template(template_dir: Path, name: str, placeholders: dict[str, Any]) -> str:
    """Load a template file and substitute {{KEY}} tokens."""
    path = template_dir / name
    if not path.exists():
        # Fall back to a minimal placeholder so the report still generates.
        return f"[Template `{name}` not found at {path}]"
    text = path.read_text(encoding="utf-8")

    def _replace(match: re.Match) -> str:
        key = match.group(1)
        val = placeholders.get(key)
        return str(val) if val is not None else f"[{key} not set]"

    return _PLACEHOLDER_RE.sub(_replace, text)


def _add_markdown_body(doc, body: str) -> None:
    """Render a simple markdown-like body string into the doc.

    Recognised forms — all template headings are SUB-SECTIONS; the writer
    code emits the top-level section H1 (e.g. "1. Introduction").
      `## subsection`  -> H2  (e.g. "## 1.1 Background")
      `### subsubsection` -> H3
      `  * bullet`      -> list item
      blank line        -> paragraph break
    """
    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        stripped = line.lstrip()
        if stripped.startswith(("* ", "- ")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:].strip())
            continue
        # Treat indented lines as a paragraph (preserve indentation as text).
        doc.add_paragraph(line.strip())


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _make_table(doc, *, headers, rows, col_widths_in=None):
    """Add a styled table with a bold header row, thin borders, centered cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(10)
        _set_cell_shading(cell, _HEADER_FILL_HEX)

    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = "" if val is None else _format_for_cell(val)
            run = p.add_run(text)
            run.font.size = Pt(10)

    # Borders
    _set_table_borders(table)

    # Optional column widths
    if col_widths_in:
        for col_idx, width_in in enumerate(col_widths_in):
            for row in table.rows:
                row.cells[col_idx].width = Inches(width_in)
    return table


def _set_table_borders(table):
    """Apply thin black borders to every cell in the table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    # Remove any existing borders block before appending the new one.
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _format_for_cell(val: Any) -> str:
    if isinstance(val, float):
        if abs(val) >= 100:
            return f"{val:.1f}"
        if abs(val) >= 1:
            return f"{val:.3f}"
        return f"{val:.4f}"
    return str(val)


# ---------------------------------------------------------------------------
# Charts (matplotlib, Agg backend)
# ---------------------------------------------------------------------------

def _chart_depth_distribution(cgr_list, out_png: Path) -> None:
    internal = [c.feature.depth_pct_wt for c in cgr_list
                if c.feature.surface is Surface.INTERNAL and c.feature.depth_pct_wt is not None]
    external = [c.feature.depth_pct_wt for c in cgr_list
                if c.feature.surface is Surface.EXTERNAL and c.feature.depth_pct_wt is not None]
    if not internal and not external:
        return
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bins = list(range(0, 105, 10))
    ax.hist(
        [internal, external], bins=bins,
        label=[f"Internal ({len(internal)})", f"External ({len(external)})"],
        color=["#2E74B5", "#ED7D31"], edgecolor="black", linewidth=0.5,
    )
    ax.set_xlabel("Depth (% WT)")
    ax.set_ylabel("Count")
    ax.set_title("Metal-Loss Defect Depth Distribution")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    ax.axvline(80, color="red", linestyle="--", alpha=0.5, label="80% repair threshold")
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)


def _chart_length_distribution(cgr_list, out_png: Path) -> None:
    lengths = [c.feature.length_mm for c in cgr_list if c.feature.length_mm is not None]
    if not lengths:
        return
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bins = np.linspace(0, max(200, max(lengths)), 21)
    ax.hist(lengths, bins=bins, color="#70AD47", edgecolor="black", linewidth=0.5)
    ax.set_xlabel("Defect Length (mm)")
    ax.set_ylabel("Count")
    ax.set_title("Metal-Loss Defect Length Distribution")
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)


def _chart_orientation_polar(cgr_list, out_png: Path) -> None:
    """Clock-position rose plot. 12:00 = top; 6:00 = bottom of pipe."""
    clocks = [c.feature.clock_decimal_hours for c in cgr_list
              if c.feature.clock_decimal_hours is not None]
    if not clocks:
        return
    # Convert clock hours -> theta in radians, with 12:00 at top (theta=π/2)
    # and clockwise increase.
    theta = [(2 * math.pi * (3 - h) / 12.0) % (2 * math.pi) for h in clocks]
    # Bin into 12 sectors (one per clock hour).
    bins = np.linspace(0, 2 * math.pi, 13)
    counts, _edges = np.histogram(theta, bins=bins)
    sector_width = 2 * math.pi / 12
    centres = bins[:-1] + sector_width / 2

    fig = plt.figure(figsize=(6.5, 6.5))
    ax = fig.add_subplot(111, projection="polar")
    ax.bar(centres, counts, width=sector_width, color="#5B9BD5",
           edgecolor="black", linewidth=0.5, alpha=0.85)
    # 12 o'clock at top
    ax.set_theta_zero_location("E")
    ax.set_theta_direction(-1)
    # Tick labels at each clock position
    hours = [12, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11]
    label_theta = [(2 * math.pi * (3 - h) / 12.0) % (2 * math.pi) for h in hours]
    ax.set_xticks(label_theta)
    ax.set_xticklabels([f"{h}" for h in hours])
    ax.set_title("Defect Orientation (Clock Position)", pad=20)
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)


def _chart_erf_acceptance(ffp_results, wt_mm: float, out_png: Path) -> None:
    """Length-vs-depth scatter colored by ERF, with 80%-WT depth line."""
    points = [(r.length_mm, r.depth_pct_wt, r.erf, r.feature_id)
              for r in ffp_results
              if r.wt_mm and abs(r.wt_mm - wt_mm) < 0.05
              and r.length_mm is not None and r.depth_pct_wt is not None]
    if not points:
        return
    Ls = [p[0] for p in points]
    Ds = [p[1] for p in points]
    Es = [p[2] for p in points]

    fig, ax = plt.subplots(figsize=(8.5, 5.5))
    sc = ax.scatter(Ls, Ds, c=Es, cmap="RdYlGn_r", alpha=0.7, s=24,
                    edgecolor="black", linewidth=0.3, vmin=0.0, vmax=max(1.0, max(Es)))
    cbar = plt.colorbar(sc, ax=ax)
    cbar.set_label("ERF")
    ax.axhline(80, color="red", linestyle="--", alpha=0.6,
               label="Depth repair criterion (80% WT)")
    ax.set_xlabel("Defect Length (mm)")
    ax.set_ylabel("Defect Depth (% WT)")
    ax.set_title(f"ERF Acceptance — WT {wt_mm} mm (n = {len(points)})")
    ax.set_xlim(left=0)
    ax.set_ylim(0, 100)
    ax.legend(loc="upper right")
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)


def _chart_repair_timeline(predictions, out_png: Path) -> None:
    """Cumulative count of repair-triggering defects by year-from-run-2."""
    triggered = [(p.repair_year_offset, p.repair_trigger) for p in predictions
                 if p.repair_year_offset is not None
                 and p.repair_trigger != TRIGGER_NONE]
    if not triggered:
        return
    max_year = max(yo for yo, _ in triggered)
    years = list(range(0, max_year + 1))
    cum_depth = [sum(1 for yo, t in triggered if t == TRIGGER_DEPTH_80 and yo <= y)
                 for y in years]
    cum_erf = [sum(1 for yo, t in triggered if t == TRIGGER_ERF_1 and yo <= y)
               for y in years]

    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    ax.step(years, cum_erf, where="post", label="ERF ≥ 1.0", color="#C00000", linewidth=2)
    ax.step(years, cum_depth, where="post", label="Depth ≥ 80% WT", color="#FFC000", linewidth=2)
    ax.set_xlabel("Years from run-2 inspection")
    ax.set_ylabel("Cumulative features requiring repair")
    ax.set_title("Predicted Repair Timeline")
    ax.grid(alpha=0.3)
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_png, dpi=130)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Table-row builders
# ---------------------------------------------------------------------------

def _shape_categorization_rows(cgr_list) -> list[list]:
    int_counter = Counter()
    ext_counter = Counter()
    for c in cgr_list:
        dim = c.feature.dimension_class
        if c.feature.surface is Surface.INTERNAL:
            int_counter[dim] += 1
        elif c.feature.surface is Surface.EXTERNAL:
            ext_counter[dim] += 1
    all_classes = sorted(set(int_counter) | set(ext_counter),
                         key=lambda d: d.value if d else "")
    total_all = sum(int_counter.values()) + sum(ext_counter.values())
    rows = []
    for dim in all_classes:
        i = int_counter.get(dim, 0)
        e = ext_counter.get(dim, 0)
        total = i + e
        pct = 100.0 * total / max(total_all, 1)
        label = dim.value if dim else "—"
        rows.append([label, i, e, total, f"{pct:.1f}%"])
    return rows


def _top_n_by_erf(ffp_results, n: int = 20) -> list[list]:
    sorted_results = sorted(ffp_results, key=lambda r: -(r.erf or 0))
    out = []
    for i, r in enumerate(sorted_results[:n], start=1):
        out.append([
            i, r.feature_id, "—", "—",
            f"{r.wt_mm:.1f}" if r.wt_mm is not None else "",
            f"{r.depth_pct_wt:.2f}" if r.depth_pct_wt is not None else "",
            f"{r.length_mm:.0f}" if r.length_mm is not None else "",
            "—",
            f"{r.sop_kgcm2:.2f}",
            f"{r.erf:.4f}",
        ])
    return out


def _top_n_by_depth(ffp_results, n: int = 20) -> list[list]:
    sorted_results = sorted(ffp_results, key=lambda r: -(r.depth_pct_wt or 0))
    out = []
    for i, r in enumerate(sorted_results[:n], start=1):
        out.append([
            i, r.feature_id, "—", "—",
            f"{r.wt_mm:.1f}" if r.wt_mm is not None else "",
            f"{r.depth_pct_wt:.2f}" if r.depth_pct_wt is not None else "",
            f"{r.length_mm:.0f}" if r.length_mm is not None else "",
            "—",
            f"{r.sop_kgcm2:.2f}",
            f"{r.erf:.4f}",
        ])
    return out


# ---------------------------------------------------------------------------
# Placeholder builder
# ---------------------------------------------------------------------------

def _build_placeholders(
    *,
    project,
    match_result,
    joint_alignment,
    cgr_results,
    ffp_results,
    repair_predictions,
    flag_report,
) -> dict[str, Any]:
    """Compute every {{KEY}} value the templates substitute."""
    ph: dict[str, Any] = {}

    pipeline = project.pipeline if project else None
    run_1 = project.run_1 if project else None
    run_2 = project.run_2 if project else None

    # Project / pipeline
    if pipeline is not None:
        ph["PIPELINE_NAME"] = pipeline.pipeline_name or ""
        ph["CLIENT_NAME"] = pipeline.client_name or ""
        ph["DIAMETER_MM"] = f"{pipeline.diameter_mm:.0f}" if pipeline.diameter_mm else "—"
        ph["DIAMETER_INCHES"] = f"{(pipeline.diameter_mm or 0) / 25.4:.1f}"
        ph["PIPELINE_LENGTH_KM"] = f"{pipeline.length_km:.1f}" if pipeline.length_km else "—"
        ph["MATERIAL_GRADE"] = pipeline.material_grade or ""
        ph["SMYS_MPA"] = f"{pipeline.smys_mpa:.0f}" if pipeline.smys_mpa else "—"
        ph["PRODUCT"] = pipeline.product or ""
        ph["SERVICE_CLASS"] = pipeline.service_class or ""
        ph["INSTALL_YEAR"] = str(pipeline.install_year or "—")
        if pipeline.maop_zones:
            # Single-zone line: render the headline value as-is.
            # Multi-zone line: keep the first-zone value in MAOP_KGCM2
            # (for any inline-template usage) but also expose a
            # zone-aware display string that lists every zone's value.
            # Without this the DOCX silently showed only zone 1, which
            # hid the actual operating-pressure spread for projects
            # like HMEL (96.7 / 84.1 / 80.6 kg/cm² across three zones).
            first = pipeline.maop_zones[0]
            ph["MAOP_KGCM2"] = f"{first.maop_kgcm2:.1f}"
            ph["DESIGN_FACTOR"] = f"{first.design_factor:.2f}"
            if len(pipeline.maop_zones) >= 2:
                # v0.3.0: zone-display label switches with mode.
                zone_word = (
                    "chainage" if getattr(pipeline, "maop_zoning_mode", "wt")
                                    == "chainage"
                    else "WT"
                )
                ph["MAOP_DISPLAY"] = (
                    " / ".join(
                        f"{z.maop_kgcm2:.1f}" for z in pipeline.maop_zones
                    )
                    + f" kg/cm² (by {zone_word} zone — see Table 2a)"
                )
                # Design factors are usually all the same; collapse if so.
                fds = {f"{z.design_factor:.2f}" for z in pipeline.maop_zones}
                ph["DESIGN_FACTOR_DISPLAY"] = (
                    f"{first.design_factor:.2f}"
                    if len(fds) == 1
                    else " / ".join(
                        f"{z.design_factor:.2f}" for z in pipeline.maop_zones
                    )
                )
            else:
                ph["MAOP_DISPLAY"] = f"{first.maop_kgcm2:.1f} kg/cm²"
                ph["DESIGN_FACTOR_DISPLAY"] = f"{first.design_factor:.2f}"
        else:
            ph["MAOP_KGCM2"] = "—"
            ph["DESIGN_FACTOR"] = "—"
            ph["MAOP_DISPLAY"] = "—"
            ph["DESIGN_FACTOR_DISPLAY"] = "—"

    if project is not None:
        ph["PROJECT_NAME"] = project.project_name or ""
        ph["REPORT_NUMBER"] = project.report_number or ""
        ph["REPORT_REVISION"] = project.report_revision or "00"
        ph["PREPARED_BY"] = project.prepared_by or ""
        ph["REVIEWED_BY"] = project.reviewed_by or ""
        ph["APPROVED_BY"] = project.approved_by or ""
        ph["REPORT_DATE"] = project.project_date.strftime("%d-%m-%Y")

    if run_1 and run_1.inspection_date:
        ph["RUN1_DATE"] = run_1.inspection_date.strftime("%d-%m-%Y")
        ph["RUN1_YEAR"] = str(run_1.inspection_date.year)
    if run_2 and run_2.inspection_date:
        ph["RUN2_DATE"] = run_2.inspection_date.strftime("%d-%m-%Y")
        ph["RUN2_YEAR"] = str(run_2.inspection_date.year)
    if run_1 and run_2 and run_1.inspection_date and run_2.inspection_date:
        delta = run_2.inspection_date - run_1.inspection_date
        ph["YEARS_BETWEEN"] = f"{delta.days / 365.25:.2f}"

    # Joint alignment
    if joint_alignment is not None:
        ph["JOINT_MATCH_COUNT"] = len(getattr(joint_alignment, "matches", []) or [])
        ph["JOINT_MATCH_RATE_PCT"] = (
            f"{(getattr(joint_alignment, 'match_rate', 0.0) * 100):.1f}"
        )
        ph["MONOTONICITY_VIOLATIONS"] = len(
            getattr(joint_alignment, "monotonicity_violations", []) or []
        )

    # Run-level feature counts
    if run_1 is not None:
        ph["RUN1_FEATURE_COUNT"] = len(run_1.features_for_assessment())
        ph["TOTAL_JOINTS_RUN1"] = len(run_1.joints)
    if run_2 is not None:
        ph["RUN2_FEATURE_COUNT"] = len(run_2.features_for_assessment())
        ph["TOTAL_JOINTS_RUN2"] = len(run_2.joints)
        ph["TOTAL_FEATURES_RUN2"] = len(run_2.features_for_assessment())

    # Match result
    if match_result is not None:
        ph["DEFECT_MATCH_COUNT"] = len(match_result.feature_matches)

    # CGR
    if cgr_results:
        internal = [c.feature_cgr_mm_yr for c in cgr_results
                    if c.feature.surface is Surface.INTERNAL]
        external = [c.feature_cgr_mm_yr for c in cgr_results
                    if c.feature.surface is Surface.EXTERNAL]
        ph["COUNT_INTERNAL"] = len(internal)
        ph["COUNT_EXTERNAL"] = len(external)
        ph["P95_CGR_INTERNAL"] = (
            f"{float(np.percentile(internal, 95)):.4f}" if len(internal) >= 2 else "—"
        )
        ph["P95_CGR_EXTERNAL"] = (
            f"{float(np.percentile(external, 95)):.4f}" if len(external) >= 2 else "—"
        )
        depths = [c.feature.depth_pct_wt for c in cgr_results
                  if c.feature.depth_pct_wt is not None]
        if depths:
            ph["MAX_DEPTH_PCT"] = f"{max(depths):.2f}"
            deepest = max(cgr_results, key=lambda c: c.feature.depth_pct_wt or 0)
            ph["MAX_DEPTH_FEATURE_ID"] = deepest.feature.anomaly_id
            ph["MAX_DEPTH_CHAINAGE_M"] = f"{deepest.feature.abs_distance_m:.1f}"
        int_pct = 100 * len(internal) / max(len(cgr_results), 1)
        ext_pct = 100 * len(external) / max(len(cgr_results), 1)
        ph["INTERNAL_COUNT"] = len(internal)
        ph["EXTERNAL_COUNT"] = len(external)
        ph["INTERNAL_PCT"] = f"{int_pct:.1f}"
        ph["EXTERNAL_PCT"] = f"{ext_pct:.1f}"
        deep = sum(1 for d in depths if d >= 50)
        ph["DEEP_FEATURE_PCT"] = f"{100 * deep / max(len(depths), 1):.1f}"

    # FFP
    if ffp_results:
        n_erf = sum(1 for r in ffp_results if r.erf >= 1.0)
        n_depth = sum(1 for r in ffp_results if r.depth_pct_wt >= 80.0)
        ph["COUNT_ERF_GE_1"] = n_erf
        ph["COUNT_DEPTH_GE_80"] = n_depth
        max_erf_result = max(ffp_results, key=lambda r: r.erf or 0)
        ph["MAX_ERF"] = f"{max_erf_result.erf:.4f}"
        ph["MAX_ERF_FEATURE_ID"] = max_erf_result.feature_id
        # Joint + chainage aren't carried on FFPResult — look them up
        # on the feature record (cgr_results carries .feature).
        _feat_by_id = {
            str(c.feature.anomaly_id): c.feature
            for c in (cgr_results or [])
        }
        _mx_feat = _feat_by_id.get(str(max_erf_result.feature_id))
        if _mx_feat is not None and _mx_feat.abs_distance_m is not None:
            ph["MAX_ERF_CHAINAGE_M"] = f"{_mx_feat.abs_distance_m:.1f}"
        else:
            ph["MAX_ERF_CHAINAGE_M"] = "—"
        if _mx_feat is not None and _mx_feat.joint_number is not None:
            ph["MAX_ERF_JOINT"] = str(_mx_feat.joint_number)
        else:
            ph["MAX_ERF_JOINT"] = "—"

    # Repair predictions
    if repair_predictions:
        triggered = [p for p in repair_predictions
                     if p.repair_trigger != TRIGGER_NONE]
        immediate = [p for p in triggered if p.repair_year_offset == 0]
        scheduled = [p for p in triggered
                     if p.repair_year_offset is not None and 0 < p.repair_year_offset <= 5]
        monitor = [p for p in repair_predictions
                   if p.repair_trigger == TRIGGER_NONE]
        ph["COUNT_REPAIR_WITHIN_HORIZON"] = len(triggered)
        ph["COUNT_IMMEDIATE"] = len(immediate)
        ph["COUNT_SCHEDULED"] = len(scheduled)
        ph["COUNT_MONITOR"] = len(monitor)
        horizon = getattr(repair_predictions[0], "horizon_years", 10) if repair_predictions else 10
        ph["HORIZON_YEARS"] = str(horizon)
        if run_2 and run_2.inspection_date:
            end = run_2.inspection_date.replace(year=run_2.inspection_date.year + horizon)
            ph["HORIZON_END_DATE"] = end.strftime("%B %Y")
    else:
        ph["COUNT_REPAIR_WITHIN_HORIZON"] = 0
        ph["COUNT_IMMEDIATE"] = 0
        ph["COUNT_SCHEDULED"] = 0
        ph["COUNT_MONITOR"] = 0
        ph["HORIZON_YEARS"] = "10"

    # CGR mode + method labels (from project config)
    cfg = (project.config or {}) if project else {}
    cgr_mode = (cfg.get("cgr") or {}).get("mode", "hybrid")
    ph["CGR_MODE"] = str(cgr_mode).upper()
    ph["FFP_METHOD"] = (cfg.get("ffp") or {}).get("primary_method", "Original")
    ph["PRIMARY_FFP_METHOD"] = ph["FFP_METHOD"]
    # POD / unmatched-depth assumption — reflects the actual cgr config.
    # The synthetic-commissioning case sets unmatched_depth_assumption_
    # pct_wt to 0.0 (brand-new pipe), so POD_PCT renders "0", not "10".
    _unmatched_pct = float(
        (cfg.get("cgr") or {}).get("unmatched_depth_assumption_pct_wt", 10.0)
    )
    ph["POD_PCT"] = f"{_unmatched_pct:g}"
    ph["REINSPECTION_INTERVAL_YEARS"] = "5"
    ph["PREPARER"] = ph.get("PREPARED_BY") or "Athena PowerTech LLP"

    # ----- Conditional narrative (v0.3.4) -------------------------------
    # When Run-1 is a synthetic commissioning baseline (an empty file —
    # no prior ILI), the standard run-to-run-comparison prose
    # (Needleman-Wunsch joint alignment, Hungarian defect matching,
    # "two inspection datasets") is dishonest. Detect that case and
    # substitute commissioning-baseline narrative. Real two-run
    # projects (Kandla, HMEL, BPCL, …) take the `else` branch and are
    # byte-for-byte unaffected.
    is_synthetic_run1 = run_1 is not None and len(run_1.features) == 0
    r1_date = ph.get("RUN1_DATE", "the commissioning date")
    r2_date = ph.get("RUN2_DATE", "the inspection date")
    r1_year = ph.get("RUN1_YEAR", "commissioning")
    r2_year = ph.get("RUN2_YEAR", "the latest run")
    yrs = ph.get("YEARS_BETWEEN", "the elapsed")
    cgr_mode_label = ph["CGR_MODE"]

    if is_synthetic_run1:
        ph["INTRO_INSPECTION_NARRATIVE"] = (
            f"One in-line inspection of the pipeline has been completed, "
            f"on {r2_date}. No prior ILI baseline exists for this line — "
            f"it was commissioned in {r1_year} and this survey is its "
            f"first inspection. For corrosion-growth estimation the "
            f"metal loss reported by this survey is conservatively "
            f"assumed to have initiated at commissioning and grown over "
            f"the intervening {yrs} years."
        )
        ph["SCOPE_COMPARISON_BULLETS"] = (
            f"  * Inventory of every metal-loss anomaly reported by the "
            f"{r2_year} baseline in-line inspection.\n"
            f"  * Conservative corrosion-growth estimation referenced "
            f"to the {r1_year} commissioning date — no earlier ILI "
            f"exists for run-to-run matching."
        )
        ph["CGR_ALIGNMENT_NARRATIVE"] = (
            f"No prior in-line inspection exists for this pipeline; the "
            f"{r2_year} survey is the baseline ILI. Conventional "
            f"run-to-run joint alignment and defect matching therefore "
            f"do not apply — there is no earlier dataset to align "
            f"against.\n\n"
            f"In the absence of a measured earlier state, every "
            f"metal-loss anomaly is conservatively assumed to have "
            f"initiated at pipeline commissioning ({r1_date}) and to "
            f"have grown linearly over the {yrs}-year period to the "
            f"{r2_date} inspection. This is the most conservative "
            f"defensible basis: each feature's entire present-day depth "
            f"is attributed to growth across that interval."
        )
        ph["CGR_DETERMINATION_NARRATIVE"] = (
            f"  * **All defects (baseline ILI)**: with no earlier "
            f"inspection, each anomaly's depth at the {r1_year} "
            f"commissioning baseline is taken as 0 % WT — a "
            f"newly-commissioned pipeline carries no metal loss. Each "
            f"feature's corrosion growth rate is therefore its full "
            f"present-day depth divided by the {yrs}-year elapsed "
            f"period: CGR = depth_run2 / {yrs}. This linear-growth-"
            f"from-commissioning assumption is conservative and yields "
            f"a strictly non-negative rate for every feature."
        )
        ph["CGR_P95_NARRATIVE"] = (
            f"To characterise the population, the 95th percentile "
            f"(P95) of the per-feature CGRs is reported separately for "
            f"internal and external defects as an upper-bound "
            f"reference. In {cgr_mode_label} mode each defect is "
            f"projected forward on its own commissioning-referenced "
            f"rate; the P95 is reported for context and is not applied "
            f"as a floor."
        )
    else:
        ph["INTRO_INSPECTION_NARRATIVE"] = (
            f"Two in-line inspections of the pipeline have been "
            f"completed: an earlier baseline run on {r1_date} and a "
            f"follow-up survey on {r2_date}, an interval of {yrs} "
            f"years."
        )
        ph["SCOPE_COMPARISON_BULLETS"] = (
            "  * Reconciliation and joint-level alignment of the two "
            "ILI datasets.\n"
            "  * Defect-level matching of metal-loss anomalies between "
            "the two runs."
        )
        ph["CGR_ALIGNMENT_NARRATIVE"] = (
            f"The two inspection datasets ({r1_year} and {r2_year}) "
            f"were aligned at the joint level using Needleman-Wunsch "
            f"sequence alignment on joint length signatures, anchored "
            f"by absolute chainage. The alignment produced "
            f"{ph.get('JOINT_MATCH_COUNT', '—')} matched joint pairs "
            f"out of a possible {ph.get('TOTAL_JOINTS_RUN2', '—')} "
            f"joints in the {r2_year} run, a match rate of "
            f"{ph.get('JOINT_MATCH_RATE_PCT', '—')}% with "
            f"{ph.get('MONOTONICITY_VIOLATIONS', '—')} chainage "
            f"reversals (zero violations indicates the alignment "
            f"preserves run-order consistency).\n\n"
            f"Within each matched joint pair, individual metal-loss "
            f"features were matched between runs using Hungarian "
            f"assignment on a cost matrix combining axial position, "
            f"clock orientation, surface side, and depth-shrinkage "
            f"plausibility. The defect-matching pipeline identified "
            f"{ph.get('DEFECT_MATCH_COUNT', '—')} cross-run pairs out "
            f"of {ph.get('RUN1_FEATURE_COUNT', '—')} {r1_year} "
            f"features and {ph.get('RUN2_FEATURE_COUNT', '—')} "
            f"{r2_year} features."
        )
        ph["CGR_DETERMINATION_NARRATIVE"] = (
            f"  * **Matched defects**: CGR = max(0, (depth_run2 − "
            f"depth_run1) / years_between). Negative growth is clamped "
            f"to zero (corrosion does not physically shrink; apparent "
            f"reduction is attributed to tool measurement "
            f"variability).\n"
            f"  * **Unmatched run-2 defects**: features reported only "
            f"in the later run are assumed to have been below the "
            f"{ph['POD_PCT']}% WT probability-of-detection threshold "
            f"at the earlier run. The CGR is then computed assuming "
            f"depth_old = {ph['POD_PCT']}% × WT, the most conservative "
            f"non-zero growth rate consistent with the tool's "
            f"detection limit."
        )
        ph["CGR_P95_NARRATIVE"] = (
            f"To capture the population-level upper bound on growth, "
            f"the 95th percentile (P95) of the feature-specific CGRs "
            f"is computed separately for internal and external "
            f"defects. In {cgr_mode_label} mode, every defect's used "
            f"growth rate is the greater of its own measured rate and "
            f"its surface P95 — features that happen to grow slowly "
            f"are not allowed to escape the population-level "
            f"expectation."
        )

    # Dent-scope reconciliation (FIX 3): if this report includes the
    # dent-strain annexure, say dents are assessed there instead of the
    # blanket "this report does not assess dents".
    _annexes = getattr(project, "report_annexures", None) or []
    _has_dent_annexure = any(
        tid == "dent_strain_b318" for tid, _letter in _annexes
    )
    if _has_dent_annexure:
        ph["SCOPE_EXCLUSIONS_NARRATIVE"] = (
            "The metal-loss fitness-for-purpose assessment in this "
            "report (Sections 2–4) does not cover dents, weld "
            "anomalies, third-party damage, or stress-corrosion "
            "cracking — those features fall outside the ASME B31G "
            "metal-loss scope. Dents are assessed separately for peak "
            "strain per ASME B31.8 §851.4.1 in the dent-strain "
            "annexure of this report. Weld anomalies, third-party "
            "damage, and stress-corrosion cracking remain outside this "
            "report's scope and are addressed by separate engineering "
            "assessments."
        )
    else:
        ph["SCOPE_EXCLUSIONS_NARRATIVE"] = (
            "This report does not assess dents, weld anomalies, "
            "third-party damage, or stress-corrosion cracking — those "
            "features fall outside the metal-loss scope and are "
            "addressed by separate engineering reports."
        )

    # Disclaimer limitations paragraph (v0.3.6 FIX A): same conditional
    # reconciliation as SCOPE_EXCLUSIONS_NARRATIVE — when the dent-strain
    # annexure is present the limitations text must not claim the report
    # "does not address … dents".
    _dent_letter = next(
        (lt for tid, lt in _annexes if tid == "dent_strain_b318"), ""
    )
    if _has_dent_annexure:
        _annex_ref = (
            f"Annexure {_dent_letter}" if _dent_letter
            else "the dent-strain annexure"
        )
        ph["LIMITATIONS_GEOMETRIC_NARRATIVE"] = (
            "This report does not address ovality, weld anomalies, "
            "third-party damage, or environmental cracking; those "
            "features fall under separate engineering analyses and the "
            "operator's general integrity-management programme. Dents "
            "are not excluded from this report — they are screened for "
            f"peak strain per ASME B31.8 §851.4.1 in {_annex_ref}."
        )
    else:
        ph["LIMITATIONS_GEOMETRIC_NARRATIVE"] = (
            "This report does not address geometrical defects (dents, "
            "ovality), weld anomalies, third-party damage, or "
            "environmental cracking. Those features fall under separate "
            "engineering analyses and the operator's general "
            "integrity-management programme."
        )

    # Conclusions narrative — auto-generated from the counts.
    n_repair = ph.get("COUNT_REPAIR_WITHIN_HORIZON", 0)
    horizon = ph.get("HORIZON_YEARS", "10")
    if n_repair == 0:
        ph["CONCLUSIONS_NARRATIVE"] = (
            f"Of the {ph.get('TOTAL_FEATURES_RUN2', '—')} metal-loss features assessed, "
            f"no defect is projected to require repair within the {horizon}-year horizon "
            "under the assumed corrosion growth rates. Continue operation at MAOP under "
            "the existing cathodic-protection programme. Re-inspect the line at the "
            "standard ILI cadence."
        )
    else:
        ph["CONCLUSIONS_NARRATIVE"] = (
            f"Of the {ph.get('TOTAL_FEATURES_RUN2', '—')} metal-loss features assessed, "
            f"{n_repair} are projected to require repair within the {horizon}-year horizon. "
            "See §4.3 for the response-category breakdown (Immediate / Scheduled / Monitor) "
            "and §4.6 for the operator's response action list."
        )

    return ph


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

def _configure_default_styles(doc) -> None:
    """Set readable defaults on the Normal + Heading styles."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for lvl, size, color in [
        ("Heading 1", 16, _HEADING_COLOR),
        ("Heading 2", 13, _HEADING_COLOR),
        ("Heading 3", 11, _HEADING_COLOR),
    ]:
        try:
            s = doc.styles[lvl]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = color
        except KeyError:
            pass
