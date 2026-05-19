"""Full-chain regression tests.

For each real project we run the *complete* pipeline once per test class
(module-scoped fixture so we don't re-run 106k features per test method)
and assert the published numbers in `examples/expected_results/*.yaml`
hold to the documented tolerances.

The two projects with full pipe-tally inputs in `/examples/`:

  1. GAIL Kandla → Samakhiali 10" LPG  (small, ~333 features)
  2. HMEL IPS-1 → IPS-2 28" crude       (large, ~106k features)

The two projects whose inputs we don't have (only the published FFP
output `.xlsx` files):

  3. GAIL Samakhiali → IP-01 (Annexure B/C/D format)
  4. HMEL COT → IPS-01 (Annexure E/F format)

For (3) and (4) we sanity-check the reference files load and have the
expected sheet structure — there's nothing else to validate without
the input pipe tallies. They're flagged as v0.2 work items.

Performance budgets (from the user's spec):

  Kandla full chain (read → annexure + DOCX):  < 30 s
  HMEL full chain (same):                       < 120 s
"""
from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from datetime import date
from pathlib import Path

import openpyxl
import pytest
import yaml
from docx import Document

from src.core.cgr import CGRCalculator
from src.core.defect_matcher import DefectMatcher
from src.core.ffp import ffp_assess
from src.core.joint_alignment import JointAligner
from src.core.repair_predictor import (
    TRIGGER_DEPTH_80, TRIGGER_ERF_1, TRIGGER_NONE, RepairPredictor,
)
from src.io.ili_reader import ILIReader
from src.models import MAOPZone, Pipeline, Project
from src.reports.annexure_writer import AnnexureWriter
from src.reports.main_report_writer import MainReportWriter
from src.validation.flag_aggregator import FlagAggregator

PROJECT_ROOT = Path(__file__).resolve().parents[1]
EXAMPLES = PROJECT_ROOT / "examples"
EXPECTED_DIR = EXAMPLES / "expected_results"
BIN_DIR = PROJECT_ROOT / "bin"


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _load_expected(name: str) -> dict:
    with (EXPECTED_DIR / name).open() as f:
        return yaml.safe_load(f)


def _run_full_chain(
    *,
    config_yaml: Path,
    out_dir: Path,
    years_between: float | None = None,
    write_docx: bool = True,
) -> dict:
    """Read → align → match → CGR → FFP → predict → aggregate → reports.

    Returns a dict of every stage's output for the tests to assert on.
    Wall-clock time is captured for the perf budgets.
    """
    t_start = time.time()

    project = Project.from_yaml(str(config_yaml))
    pipeline = project.pipeline
    runs_cfg = project.config.get("runs") or {}
    run1_rel = runs_cfg.get("run_1", {}).get("file_path")
    run2_rel = runs_cfg.get("run_2", {}).get("file_path")
    run1_path = PROJECT_ROOT / run1_rel
    run2_path = PROJECT_ROOT / run2_rel
    assert run1_path.exists(), f"run-1 missing: {run1_path}"
    assert run2_path.exists(), f"run-2 missing: {run2_path}"

    reader = ILIReader()
    run1 = reader.read(str(run1_path), run_id="run_1")
    if project.run_1.inspection_date:
        run1.inspection_date = project.run_1.inspection_date
    run2 = reader.read(str(run2_path), run_id="run_2")
    if project.run_2.inspection_date:
        run2.inspection_date = project.run_2.inspection_date
    project.run_1 = run1
    project.run_2 = run2

    if years_between is None:
        years_between = (
            (run2.inspection_date - run1.inspection_date).days / 365.25
        )

    ja = JointAligner().align(run1, run2)
    mr = DefectMatcher().match(run1, run2, ja.matches)

    cgr_mode = (project.config.get("cgr") or {}).get("mode", "hybrid")
    cgrs = CGRCalculator({"mode": cgr_mode}).compute(mr, years_between=years_between)

    primary = (project.config.get("ffp") or {}).get("primary_method", "B31G_Original")
    ffps_by_id = {}
    for c in cgrs:
        try:
            fl = ffp_assess(c.feature, pipeline, config={"primary_method": primary})
            ffps_by_id[c.feature.anomaly_id] = next(
                (f for f in fl if f.is_controlling), fl[0]
            )
        except ValueError:
            continue

    preds = RepairPredictor().predict(
        cgrs, ffps_by_id, pipeline,
        run2_inspection_date=run2.inspection_date,
    )
    flag_report = FlagAggregator().aggregate(
        run1=run1, run2=run2, joint_alignment=ja, match_result=mr,
        cgr_results=cgrs, ffp_results=list(ffps_by_id.values()),
        predictions=preds,
    )

    out_dir.mkdir(parents=True, exist_ok=True)
    annex_path = out_dir / "annexure.xlsx"
    AnnexureWriter().write(
        cgr_results=cgrs, ffp_results=list(ffps_by_id.values()),
        repair_predictions=preds, flag_report=flag_report,
        project=project, pipeline=pipeline,
        output_path=str(annex_path),
        format="E_F",
    )
    docx_path = out_dir / "report.docx" if write_docx else None
    if write_docx:
        MainReportWriter().write(
            project=project, match_result=mr, joint_alignment=ja,
            cgr_results=cgrs, ffp_results=list(ffps_by_id.values()),
            repair_predictions=preds, flag_report=flag_report,
            output_path=str(docx_path),
        )

    elapsed = time.time() - t_start

    return {
        "project": project, "pipeline": pipeline,
        "run1": run1, "run2": run2,
        "joint_alignment": ja, "match_result": mr,
        "cgr_results": cgrs, "ffp_results": list(ffps_by_id.values()),
        "ffps_by_id": ffps_by_id,
        "predictions": preds, "flag_report": flag_report,
        "annexure_path": annex_path, "docx_path": docx_path,
        "elapsed_s": elapsed, "years_between": years_between,
    }


def _within_pct(actual: float, expected: float, tol_pct: float) -> bool:
    """`|actual - expected| / |expected| <= tol_pct`. Returns True for
    expected==0 only when actual==0 exactly."""
    if expected == 0:
        return actual == 0
    return abs(actual - expected) / abs(expected) <= tol_pct


# ---------------------------------------------------------------------------
# Project 1 — GAIL Kandla → Samakhiali 10" LPG
# ---------------------------------------------------------------------------

class TestKandlaRegression:
    """End-to-end on the Kandla pair against `kandla_samakhiali.yaml`."""

    @pytest.fixture(scope="class")
    def chain(self, tmp_path_factory):
        out = tmp_path_factory.mktemp("kandla_reg")
        return _run_full_chain(
            config_yaml=EXAMPLES / "kandla_project.yaml",
            out_dir=out,
            years_between=4.25,        # matches the published report
        )

    @pytest.fixture(scope="class")
    def expected(self):
        return _load_expected("kandla_samakhiali.yaml")

    # ---- Counts ----

    def test_total_features(self, chain, expected):
        assert len(chain["cgr_results"]) == expected["total_features"]

    def test_matched_defect_count_meets_floor(self, chain, expected):
        n = len(chain["match_result"].feature_matches)
        assert n >= expected["matched_features_min"], (
            f"matched count {n} below floor {expected['matched_features_min']}"
        )

    # ---- Canonical defect #125 ----

    def test_feature_125_cgr_exact(self, chain, expected):
        f125 = next((r for r in chain["cgr_results"]
                     if r.feature.anomaly_id == "125"), None)
        assert f125 is not None, "feature #125 missing from CGR results"
        e = expected["feature_125"]
        assert abs(f125.feature_cgr_mm_yr - e["cgr_mm_yr"]) <= e["cgr_tolerance"], (
            f"#125 CGR {f125.feature_cgr_mm_yr:.5f} vs published "
            f"{e['cgr_mm_yr']:.5f}"
        )

    def test_feature_125_psafe(self, chain, expected):
        ffp = chain["ffps_by_id"].get("125")
        assert ffp is not None
        e = expected["feature_125"]
        assert _within_pct(ffp.sop_kgcm2, e["psafe_kgcm2"], e["psafe_tolerance_pct"]), (
            f"#125 Psafe {ffp.sop_kgcm2:.2f} kg/cm² vs published {e['psafe_kgcm2']}"
        )

    def test_feature_125_erf_today(self, chain, expected):
        ffp = chain["ffps_by_id"].get("125")
        e = expected["feature_125"]
        assert _within_pct(ffp.erf, e["erf_at_year_0"], e["erf_tolerance_pct"]), (
            f"#125 ERF (year 0) {ffp.erf:.4f} vs published {e['erf_at_year_0']}"
        )

    def test_feature_125_erf_year10(self, chain, expected):
        pred = next((p for p in chain["predictions"]
                     if p.feature.anomaly_id == "125"), None)
        assert pred is not None
        e = expected["feature_125"]
        assert _within_pct(
            pred.final_erf, e["erf_at_year_10"], e["erf_year10_tolerance_pct"]
        ), (
            f"#125 ERF (year 10) {pred.final_erf:.4f} vs published "
            f"{e['erf_at_year_10']}"
        )

    # ---- Population stats ----

    def test_internal_p95_cgr(self, chain, expected):
        from src.models import Surface
        internal = [c.feature_cgr_mm_yr for c in chain["cgr_results"]
                    if c.feature.surface is Surface.INTERNAL]
        import numpy as np
        actual = float(np.percentile(internal, 95))
        assert _within_pct(
            actual, expected["internal_p95_cgr_mm_yr"],
            expected["internal_p95_tolerance_pct"]
        ), f"internal P95 {actual:.4f} vs published {expected['internal_p95_cgr_mm_yr']}"

    def test_external_p95_cgr(self, chain, expected):
        from src.models import Surface
        external = [c.feature_cgr_mm_yr for c in chain["cgr_results"]
                    if c.feature.surface is Surface.EXTERNAL]
        import numpy as np
        actual = float(np.percentile(external, 95))
        assert _within_pct(
            actual, expected["external_p95_cgr_mm_yr"],
            expected["external_p95_tolerance_pct"]
        ), f"external P95 {actual:.4f} vs published {expected['external_p95_cgr_mm_yr']}"

    # ---- Repair prediction ----

    def test_all_features_none_within_horizon(self, chain, expected):
        preds = chain["predictions"]
        triggered = [p for p in preds if p.repair_trigger != TRIGGER_NONE]
        assert len(triggered) == expected["expected_repairs_in_10yr"], (
            f"{len(triggered)} feature(s) flagged for repair; expected "
            f"{expected['expected_repairs_in_10yr']}"
        )

    # ---- QA ----

    def test_qa_has_critical_false(self, chain, expected):
        assert chain["flag_report"].has_critical is expected["expected_has_critical"]

    # ---- Annexure ----

    def test_annexure_e_row_count(self, chain, expected):
        wb = openpyxl.load_workbook(chain["annexure_path"])
        assert "Annexure E" in wb.sheetnames
        ws = wb["Annexure E"]
        # 3 header rows + N data rows
        n_data = ws.max_row - 3
        assert n_data >= expected["annexure_e_data_rows_min"], (
            f"Annexure E has {n_data} data rows; expected ≥ "
            f"{expected['annexure_e_data_rows_min']}"
        )

    # ---- DOCX verdict ----

    def test_docx_verdict_phrase(self, chain, expected):
        doc = Document(str(chain["docx_path"]))
        body = "\n".join(p.text for p in doc.paragraphs)
        assert expected["verdict_phrase"] in body, (
            f"verdict phrase {expected['verdict_phrase']!r} not in DOCX body"
        )

    # ---- Perf ----

    def test_runtime_under_budget(self, chain, expected):
        budget = expected["runtime_budget_s"]
        assert chain["elapsed_s"] < budget, (
            f"full chain took {chain['elapsed_s']:.1f}s, budget {budget}s"
        )


# ---------------------------------------------------------------------------
# Project 2 — HMEL IPS-1 → IPS-2 28" crude
# ---------------------------------------------------------------------------

class TestHMELRegression:
    """End-to-end on the HMEL pair against `hmel_ips1_ips2.yaml`."""

    @pytest.fixture(scope="class")
    def chain(self, tmp_path_factory):
        out = tmp_path_factory.mktemp("hmel_reg")
        return _run_full_chain(
            config_yaml=EXAMPLES / "hmel_ips1_ips2_project.yaml",
            out_dir=out,
            years_between=5.42,          # May 2019 → Oct 2025
        )

    @pytest.fixture(scope="class")
    def expected(self):
        return _load_expected("hmel_ips1_ips2.yaml")

    # ---- Counts ----

    def test_matched_defect_count_meets_floor(self, chain, expected):
        n = len(chain["match_result"].feature_matches)
        assert n >= expected["matched_features_min"], (
            f"matched count {n} below floor {expected['matched_features_min']}"
        )

    def test_features_for_assessment_count(self, chain, expected):
        n = len(chain["cgr_results"])
        assert _within_pct(
            n, expected["features_for_assessment"],
            expected["features_for_assessment_tolerance_pct"]
        ), (
            f"features-for-assessment {n} vs published "
            f"{expected['features_for_assessment']}"
        )

    # ---- Canonical defect #209581 ----

    def test_feature_209581_psafe(self, chain, expected):
        ffp = chain["ffps_by_id"].get("209581")
        assert ffp is not None, "feature #209581 missing"
        e = expected["feature_209581"]
        # Absolute-tolerance check now that the published value is an
        # exact reconciliation (psafe_tolerance is 0.1 kg/cm²).
        tol = e.get("psafe_tolerance", 0.1)
        assert abs(ffp.sop_kgcm2 - e["psafe_kgcm2"]) <= tol, (
            f"#209581 Psafe {ffp.sop_kgcm2:.2f} kg/cm² vs published "
            f"{e['psafe_kgcm2']} ± {tol}"
        )

    def test_feature_209581_erf_today(self, chain, expected):
        """With the corrected MAOP-zone mapping (#209581 sits in zone 1
        at MAOP=80.6, NOT the old zone-3 inversion), the tool now hits
        the published ERF exactly. Tolerance is ±0.005 to absorb
        rounding in Psafe.
        """
        ffp = chain["ffps_by_id"].get("209581")
        e = expected["feature_209581"]
        tol = e.get("erf_tolerance", 0.005)
        assert abs(ffp.erf - e["erf_today"]) <= tol, (
            f"#209581 ERF {ffp.erf:.3f} vs published {e['erf_today']} "
            f"± {tol}"
        )

    def test_feature_209581_maop_zone_assignment(self, chain, expected):
        """#209581 (WT=8.7 mm) must land in zone 1 at MAOP=80.6 kg/cm²,
        not the historical WT-inverted assignment that put it at 96.7."""
        ffp = chain["ffps_by_id"].get("209581")
        e = expected["feature_209581"]
        assert abs(ffp.maop_kgcm2 - e["maop_kgcm2_used"]) <= 0.5, (
            f"#209581 used MAOP={ffp.maop_kgcm2}, expected "
            f"{e['maop_kgcm2_used']} (zone 1)"
        )

    # ---- Population ERF stats ----

    def test_features_with_erf_above_1(self, chain, expected):
        """Published exec summary: 7 features with ERF > 1.0 (using the
        operator's MAOP-zone mapping, which the tool now matches)."""
        n_erf_above_1 = sum(1 for r in chain["ffp_results"] if r.erf >= 1.0)
        target = expected["features_with_erf_above_1"]
        tol = expected.get("features_with_erf_above_1_tolerance", 2)
        assert abs(n_erf_above_1 - target) <= tol, (
            f"got {n_erf_above_1} features with ERF≥1.0; published "
            f"target is {target} ± {tol}"
        )

    def test_top_erf_feature_is_209581(self, chain, expected):
        """The single highest-ERF feature must be #209581 — confirms the
        zone mapping is applied uniformly across all features, not just
        the spot-check one."""
        if not chain["ffp_results"]:
            return
        top = max(chain["ffp_results"], key=lambda r: r.erf)
        assert top.feature_id == expected["top_erf_feature_id"], (
            f"top-ERF feature is {top.feature_id} (ERF={top.erf:.3f}); "
            f"expected {expected['top_erf_feature_id']}"
        )

    # ---- Annexure ----

    def test_annexure_loads_with_both_sheets(self, chain):
        wb = openpyxl.load_workbook(chain["annexure_path"], read_only=True)
        assert "Annexure E" in wb.sheetnames
        assert "Annexure F" in wb.sheetnames
        wb.close()

    # ---- DOCX ----

    def test_docx_loads(self, chain):
        doc = Document(str(chain["docx_path"]))
        h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "1. Introduction" in h1
        # HMEL has features above ERF=1 → exec summary should NOT say "No defects require repair"
        body = "\n".join(p.text for p in doc.paragraphs)
        assert "Defects flagged for action" in body or \
               "require repair within the" in body

    # ---- Perf ----

    def test_runtime_under_budget(self, chain, expected):
        budget = expected["runtime_budget_s"]
        assert chain["elapsed_s"] < budget, (
            f"full chain took {chain['elapsed_s']:.1f}s, budget {budget}s"
        )


# ---------------------------------------------------------------------------
# Projects 3 & 4 — output-only references (no inputs in repo)
# ---------------------------------------------------------------------------

class TestPublishedReferenceFilesLoad:
    """For Samakhiali → IP-01 and COT → IPS-01 we only have the published
    annexure files, not the input pipe tallies. Sanity-check the
    references load — full regression deferred to v0.2 when the input
    files arrive."""

    def test_samakhiali_bcd_loads(self):
        """The .xls extension is misleading — file is XLSX-by-magic-bytes.
        Copy to .xlsx in tempdir to satisfy openpyxl's extension check."""
        src = EXAMPLES / "IPS__Samakhiali_to_IP_01_-_Annexure_B__C__D_FFP_.xls"
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            shutil.copy(src, tmp.name)
            tmp_path = Path(tmp.name)
        try:
            wb = openpyxl.load_workbook(tmp_path, read_only=True)
            sheet_names = wb.sheetnames
            assert "Annexure B" in sheet_names
            assert "Annexure C" in sheet_names
            assert "Annexure D" in sheet_names
            wb.close()
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_hmel_cot_to_ips_ef_loads(self):
        ref = EXAMPLES / "FFP_Report_COT_to_IPS_1_Annexure_E__F.xlsx"
        wb = openpyxl.load_workbook(ref, read_only=True)
        assert "Annexure E" in wb.sheetnames
        assert "Annexure F" in wb.sheetnames
        wb.close()


# ---------------------------------------------------------------------------
# CLI runner — bin/run_pipeline.py
# ---------------------------------------------------------------------------

class TestCLIRunner:
    """Verify `python bin/run_pipeline.py` produces the same outputs as the
    in-process pipeline. Exit code policy: 0 on clean, 1 on has_critical."""

    def test_cli_kandla_exits_zero(self, tmp_path):
        out = tmp_path / "out"
        result = subprocess.run(
            [sys.executable, str(BIN_DIR / "run_pipeline.py"),
             "--config", str(EXAMPLES / "kandla_project.yaml"),
             "--output-dir", str(out),
             "--quiet"],
            cwd=str(PROJECT_ROOT),
            capture_output=True,
        )
        assert result.returncode == 0, (
            f"CLI exit {result.returncode}; stderr:\n"
            f"{result.stderr.decode('utf-8', errors='replace')[:1000]}"
        )
        # Outputs should exist.
        annex = list(out.glob("*_annexure.xlsx"))
        docx = list(out.glob("*_report.docx"))
        assert annex, "no annexure produced"
        assert docx, "no DOCX produced"

    def test_cli_no_docx_flag_skips_docx(self, tmp_path):
        out = tmp_path / "out"
        result = subprocess.run(
            [sys.executable, str(BIN_DIR / "run_pipeline.py"),
             "--config", str(EXAMPLES / "kandla_project.yaml"),
             "--output-dir", str(out),
             "--no-docx", "--quiet"],
            cwd=str(PROJECT_ROOT),
            capture_output=True,
        )
        assert result.returncode == 0, result.stderr[:500]
        assert list(out.glob("*_annexure.xlsx"))
        assert not list(out.glob("*_report.docx"))

    def test_cli_missing_config_returns_2(self, tmp_path):
        out = tmp_path / "out"
        result = subprocess.run(
            [sys.executable, str(BIN_DIR / "run_pipeline.py"),
             "--config", str(tmp_path / "does_not_exist.yaml"),
             "--output-dir", str(out),
             "--quiet"],
            cwd=str(PROJECT_ROOT),
            capture_output=True,
        )
        assert result.returncode == 2
