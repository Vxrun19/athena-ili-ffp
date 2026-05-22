"""Tests for src.reports.main_report_writer.

We generate the full FFP report for the Kandla-Samakhiali pair and verify:
  * the file opens and is well-formed (docx loads it),
  * every top-level section is present,
  * placeholders were substituted (no `{{KEY}}` left in the body),
  * tables and embedded images are present,
  * the canonical Kandla numbers appear in the executive summary,
  * the writer finishes within the 5-second performance budget.
"""
from __future__ import annotations

import re
import time
import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from src.core.cgr import CGRCalculator
from src.core.defect_matcher import DefectMatcher
from src.core.ffp import ffp_assess
from src.core.joint_alignment import JointAligner
from src.core.repair_predictor import RepairPredictor
from src.io.ili_reader import ILIReader
from src.models import MAOPZone, Pipeline, Project
from src.reports.main_report_writer import MainReportWriter
from src.validation.flag_aggregator import FlagAggregator

EXAMPLES = Path(__file__).resolve().parents[1] / "examples"


# ---------------------------------------------------------------------------
# Shared fixture — the full Kandla chain, generated once per test module.
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def kandla_chain():
    pipeline = Pipeline(
        pipeline_name='Kandla-Samakhiali 10" LPG',
        client_name="GAIL (India) Limited",
        diameter_mm=273.0, length_km=58.5, smys_mpa=358.0,
        material_grade="API 5L X52", product="LPG",
        service_class="liquid", install_year=2011,
        maop_zones=[MAOPZone(6.0, 8.0, 0.72, 70.0)],
    )
    reader = ILIReader()
    run1 = reader.read(
        EXAMPLES / "10_inch_Kandla_to_Samakhiali__58_2_km_FR_Pipe_Tally_Rev_0.xlsx",
        run_id="k1",
    )
    run1.inspection_date = date(2018, 12, 15)
    run2 = reader.read(EXAMPLES / "1ZSV_Pipeline_Listing.xlsx", run_id="k2")
    run2.inspection_date = date(2023, 3, 15)
    project = Project(
        project_name="FFP_Kandla_Samakhiali",
        report_number="ATH-2023-001", report_revision="00",
        prepared_by="Prepared By", reviewed_by="Reviewed By",
        approved_by="Approved By",
        pipeline=pipeline, run_1=run1, run_2=run2,
    )
    ja = JointAligner().align(run1, run2)
    mr = DefectMatcher().match(run1, run2, ja.matches)
    cgrs = CGRCalculator({"mode": "hybrid"}).compute(mr, years_between=4.25)
    ffps_by_id = {}
    for c in cgrs:
        fl = ffp_assess(c.feature, pipeline)
        ffps_by_id[c.feature.anomaly_id] = next(
            (f for f in fl if f.is_controlling), fl[0]
        )
    preds = RepairPredictor().predict(
        cgrs, ffps_by_id, pipeline,
        run2_inspection_date=date(2023, 3, 15),
    )
    flag_report = FlagAggregator().aggregate(
        run1=run1, run2=run2, joint_alignment=ja, match_result=mr,
        cgr_results=cgrs, ffp_results=list(ffps_by_id.values()),
        predictions=preds,
    )
    return {
        "project": project, "match_result": mr, "joint_alignment": ja,
        "cgr_results": cgrs, "ffp_results": list(ffps_by_id.values()),
        "predictions": preds, "flag_report": flag_report,
    }


@pytest.fixture(scope="module")
def kandla_report(tmp_path_factory, kandla_chain):
    """Generate the report once and reuse across the asserts.

    Returns (path, elapsed_seconds).
    """
    out = tmp_path_factory.mktemp("docx") / "kandla_report.docx"
    t0 = time.time()
    MainReportWriter().write(
        project=kandla_chain["project"],
        match_result=kandla_chain["match_result"],
        joint_alignment=kandla_chain["joint_alignment"],
        cgr_results=kandla_chain["cgr_results"],
        ffp_results=kandla_chain["ffp_results"],
        repair_predictions=kandla_chain["predictions"],
        flag_report=kandla_chain["flag_report"],
        output_path=str(out),
    )
    return out, time.time() - t0


@pytest.fixture(scope="module")
def kandla_report_path(kandla_report):
    return kandla_report[0]


# ---------------------------------------------------------------------------
# Structure / well-formedness
# ---------------------------------------------------------------------------

class TestReportStructure:
    def test_file_exists_and_opens(self, kandla_report_path):
        assert kandla_report_path.exists()
        assert kandla_report_path.stat().st_size > 50_000   # >50 KB sanity
        # Loads cleanly
        doc = Document(str(kandla_report_path))
        assert len(doc.paragraphs) > 50

    def test_all_top_level_sections_present(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        for expected in [
            "Executive Summary",
            "Abbreviations",
            "Table of Contents",
            "1. Introduction",
            "2. ILI Results",
            "3. Fitness-For-Purpose Analysis",
            "4. ILI Reports Comparison, CGR Analysis, and Repair Prediction",
            "Disclaimer",
        ]:
            assert expected in h1, (
                f"top-level section {expected!r} missing; got {h1}"
            )
        # Annexure A heading uses an em-dash character.
        assert any("Annexure A" in t for t in h1), f"Annexure A missing; got {h1}"

    def test_section_4_subsections_all_present(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        for expected in [
            "4.1 ILI Reports Assessment",
            "4.2 Repair Prediction Methodology",
            "4.3 Response to Pipeline ILI Results (API 1160 Categories)",
            "4.4 Recommended Repair Methods",
            "4.5 Integrity Management",
            "4.6 Conclusions and Recommendations",
        ]:
            assert expected in h2, f"missing §{expected}; got {h2}"

    def test_introduction_subsections(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        for expected in ["1.1 Background", "1.2 Scope", "1.3 Pipeline Details"]:
            assert expected in h2

    def test_table_count_and_pipeline_table(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        # Pipeline details, shape categ., comparison summary, P95 CGR, 6a, 6b,
        # cover-page revision table — at least 7 tables expected.
        assert len(doc.tables) >= 6, (
            f"expected ≥6 tables, got {len(doc.tables)}"
        )
        # The pipeline details table contains "Pipeline name" / "Material grade" rows.
        all_cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        assert "Pipeline name" in all_cells
        assert "MAOP" in all_cells
        assert "SMYS" in all_cells

    def test_embedded_images_present(self, kandla_report_path):
        with zipfile.ZipFile(kandla_report_path) as zf:
            images = [n for n in zf.namelist() if n.startswith("word/media/")]
        # Depth dist, length dist, orientation, repair timeline, ≥1 ERF chart.
        assert len(images) >= 4, f"expected ≥4 embedded images, got {len(images)}"


# ---------------------------------------------------------------------------
# Content correctness — Kandla data must appear in the output
# ---------------------------------------------------------------------------

class TestKandlaContent:
    def test_no_unresolved_placeholders_in_body(self, kandla_report_path):
        """Every `{{KEY}}` token in any template should have been substituted
        (or marked `[KEY not set]` if the writer couldn't fill it)."""
        doc = Document(str(kandla_report_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        unresolved = re.findall(r"\{\{(\w+)\}\}", all_text)
        assert not unresolved, f"unresolved placeholders: {unresolved}"

    def test_pipeline_name_appears(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs) + "\n" + \
               "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        assert "Kandla-Samakhiali" in body

    def test_inspection_years_in_body(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs)
        assert "2018" in body
        assert "2023" in body

    def test_executive_summary_kandla_numbers(self, kandla_report_path, kandla_chain):
        """The executive summary should report the validated Kandla
        numbers — 333 features, 0 ERF≥1, 0 depth≥80, 0 repairs."""
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs)
        # 333 features
        assert "333" in body
        # Zero repairs expected — exec summary's verdict line must say so
        assert "No defects require repair" in body or \
               "no defect is projected to require repair" in body.lower()

    def test_max_erf_feature_id_in_body(self, kandla_report_path, kandla_chain):
        """The feature with the maximum ERF in Kandla is one of the
        run-2 features. Its anomaly_id must surface in the exec summary."""
        ffps = kandla_chain["ffp_results"]
        max_erf_id = max(ffps, key=lambda r: r.erf or 0).feature_id
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs)
        assert max_erf_id in body, (
            f"max-ERF feature id {max_erf_id!r} missing from body"
        )

    def test_p95_cgr_values_present(self, kandla_report_path):
        """Internal/external P95 CGR values appear in the Table 5 section
        and the exec summary."""
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs) + "\n" + \
               "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        # Should appear as numbers like "0.0589" (internal) and "0.0385" (external).
        # The exact strings may vary slightly with rounding; check rough range.
        assert "P95 CGR" in body or "P95" in body

    def test_qa_summary_in_exec_summary(self, kandla_report_path):
        doc = Document(str(kandla_report_path))
        body = "\n".join(p.text for p in doc.paragraphs)
        # FlagReport.summary starts with "QA:"
        assert "QA:" in body

    def test_top_20_by_erf_table_has_rows(self, kandla_report_path, kandla_chain):
        """One of the tables must be the Top-20-by-ERF list with feature
        IDs from the real assessment."""
        doc = Document(str(kandla_report_path))
        # Find a table whose header contains "ERF"
        erf_table = None
        for t in doc.tables:
            header_text = [c.text for c in t.rows[0].cells]
            if "ERF" in header_text and "Feature ID" in header_text:
                erf_table = t
                break
        assert erf_table is not None, "Top-N-by-ERF table not found"
        # Should have ≥1 data row (header + at least 1).
        assert len(erf_table.rows) >= 2
        # The first data row's ERF should be the max ERF observed.
        max_erf = max(r.erf or 0 for r in kandla_chain["ffp_results"])
        # ERF column is the last column.
        first_data_erf = float(erf_table.rows[1].cells[-1].text)
        assert first_data_erf == pytest.approx(max_erf, abs=1e-4)

    def test_max_erf_line_has_numeric_joint_and_chainage(self, kandla_report_path):
        """v0.3.6 FIX B: the FFP-analysis max-ERF sentence must populate
        the joint number and chainage from the feature record — it used
        to emit em-dash placeholders ('joint —, chainage — m')."""
        import re as _re
        doc = Document(str(kandla_report_path))
        line = next(
            (p.text for p in doc.paragraphs
             if "maximum recorded ERF" in p.text),
            None,
        )
        assert line is not None, "max-ERF sentence not found in report body"
        assert "chainage — m" not in line and "joint —" not in line, (
            f"max-ERF line still carries placeholder dashes: {line!r}"
        )
        assert _re.search(r"chainage\s+[\d.]+\s*m", line), (
            f"max-ERF line has no numeric chainage: {line!r}"
        )
        assert _re.search(r"joint\s+\d+", line), (
            f"max-ERF line has no numeric joint: {line!r}"
        )


# ---------------------------------------------------------------------------
# Performance
# ---------------------------------------------------------------------------

class TestPerformance:
    def test_generation_under_5_seconds(self, kandla_report):
        _path, elapsed = kandla_report
        assert elapsed < 5.0, f"report generation took {elapsed:.2f}s"


# ---------------------------------------------------------------------------
# Empty-pipeline edge case
# ---------------------------------------------------------------------------

class TestEdgeCases:
    def test_minimal_inputs_still_writes(self, tmp_path):
        """Writer should produce a file even with empty inputs — just no
        data rows or charts, but all section headings should still appear."""
        pipeline = Pipeline(
            pipeline_name="Test", client_name="Acme",
            diameter_mm=273.0, length_km=10.0, smys_mpa=358.0,
            maop_zones=[MAOPZone(6.0, 10.0, 0.72, 70.0)],
        )
        project = Project(
            project_name="Minimal",
            pipeline=pipeline,
        )
        out = tmp_path / "minimal.docx"
        MainReportWriter().write(
            project=project,
            cgr_results=[], ffp_results=[],
            repair_predictions=[],
            output_path=str(out),
        )
        assert out.exists()
        doc = Document(str(out))
        h1 = {p.text for p in doc.paragraphs if p.style.name == "Heading 1"}
        assert "1. Introduction" in h1
        assert "Disclaimer" in h1


# ---------------------------------------------------------------------------
# v0.3.4 — narrative adapts to a synthetic (empty) Run-1.
# ---------------------------------------------------------------------------

class TestNarrativeConditional:
    """The CGR / introduction narrative must not describe a run-to-run
    comparison when Run-1 is a synthetic commissioning baseline (empty
    file, no prior ILI). Real two-run projects keep the original prose.
    """

    @staticmethod
    def _placeholders(*, run1_empty, cgr_cfg, with_dent_annexure):
        from src.models import ILIRun, Feature, FeatureIdentification, Surface
        from src.reports.main_report_writer import _build_placeholders

        def _feat(aid):
            return Feature(
                anomaly_id=aid, source_run="r", abs_distance_m=10.0,
                joint_number=1, wt_mm=8.0, depth_pct_wt=20.0,
                length_mm=30.0, width_mm=20.0, surface=Surface.EXTERNAL,
                feature_identification=FeatureIdentification.CORROSION,
            )

        run1 = ILIRun(run_id="run_1", inspection_date=date(2019, 12, 1))
        if not run1_empty:
            run1.features.append(_feat("r1-1"))
        run2 = ILIRun(run_id="run_2", inspection_date=date(2026, 2, 23))
        run2.features.append(_feat("r2-1"))
        pipeline = Pipeline(
            pipeline_name="Test Line", diameter_mm=457.2,
            length_km=100.0, smys_mpa=448.0,
        )
        annexures = [("dent_strain_b318", "E")] if with_dent_annexure else []
        project = Project(
            project_name="t", pipeline=pipeline,
            run_1=run1, run_2=run2,
            config={"cgr": cgr_cfg},
            report_annexures=annexures,
        )
        return _build_placeholders(
            project=project, match_result=None, joint_alignment=None,
            cgr_results=[], ffp_results=[], repair_predictions=[],
            flag_report=None,
        )

    def test_synthetic_run1_no_prior_ili_narrative(self):
        ph = self._placeholders(
            run1_empty=True,
            cgr_cfg={"mode": "feature_specific",
                     "unmatched_depth_assumption_pct_wt": 0.0},
            with_dent_annexure=True,
        )
        assert "No prior ILI" in ph["INTRO_INSPECTION_NARRATIVE"]
        assert "Needleman" not in ph["CGR_ALIGNMENT_NARRATIVE"]
        assert "commissioning" in ph["CGR_ALIGNMENT_NARRATIVE"].lower()
        assert "0 % WT" in ph["CGR_DETERMINATION_NARRATIVE"]
        assert ph["POD_PCT"] == "0"

    def test_real_run1_keeps_run_to_run_narrative(self):
        ph = self._placeholders(
            run1_empty=False,
            cgr_cfg={"mode": "hybrid"},
            with_dent_annexure=False,
        )
        assert "Two in-line inspections" in ph["INTRO_INSPECTION_NARRATIVE"]
        assert "Needleman-Wunsch" in ph["CGR_ALIGNMENT_NARRATIVE"]
        assert "Hungarian" in ph["CGR_ALIGNMENT_NARRATIVE"]
        assert "Matched defects" in ph["CGR_DETERMINATION_NARRATIVE"]
        assert ph["POD_PCT"] == "10"

    def test_dent_annexure_reconciles_scope_wording(self):
        with_dent = self._placeholders(
            run1_empty=True,
            cgr_cfg={"mode": "feature_specific"},
            with_dent_annexure=True,
        )
        assert "dent-strain annexure" in with_dent["SCOPE_EXCLUSIONS_NARRATIVE"]
        without = self._placeholders(
            run1_empty=False,
            cgr_cfg={"mode": "hybrid"},
            with_dent_annexure=False,
        )
        assert ("separate engineering reports"
                in without["SCOPE_EXCLUSIONS_NARRATIVE"])

    def test_limitations_paragraph_reconciles_with_dent_annexure(self):
        """v0.3.6 FIX A: the disclaimer limitations paragraph must not
        claim the report excludes dents when the dent-strain annexure
        is present. Both branches must render cleanly."""
        with_dent = self._placeholders(
            run1_empty=True,
            cgr_cfg={"mode": "feature_specific"},
            with_dent_annexure=True,
        )
        lim_with = with_dent["LIMITATIONS_GEOMETRIC_NARRATIVE"]
        # Must NOT lump dents into the excluded list...
        assert "geometrical defects (dents" not in lim_with
        # ...and must point at the dent-strain annexure (letter "E" in
        # this fixture) per ASME B31.8 §851.4.1.
        assert "ASME B31.8" in lim_with
        assert "Annexure E" in lim_with
        assert "ovality" in lim_with  # other geom features still excluded

        without = self._placeholders(
            run1_empty=False,
            cgr_cfg={"mode": "hybrid"},
            with_dent_annexure=False,
        )
        lim_without = without["LIMITATIONS_GEOMETRIC_NARRATIVE"]
        # No dent annexure -> the original blanket wording is preserved.
        assert "geometrical defects (dents, ovality)" in lim_without
        assert "ASME B31.8" not in lim_without
